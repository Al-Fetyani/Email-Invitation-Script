import pandas as pd
import smtplib
from email.utils import formataddr
from docx import Document
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import threading
import time
import argparse
import getpass


def fill_invitation(template_path, data):
    doc = Document(template_path)

    for p in doc.paragraphs:
        for key, value in data.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, value)
    doc_content = "\n".join([p.text for p in doc.paragraphs])

    return doc_content


def send_email(to_email, subject, username, password, body):
    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"] = formataddr(("Al-Fetyani", username))
    msg.attach(MIMEText(body, "plain", "utf-8"))

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(username, password)
            server.sendmail(username, to_email, msg.as_string())
        print(f"Email sent to: {to_email}")
    except Exception as e:
        print(f"Error sending email to {to_email}: {e}")


def send_emails(emails, subject, username, password):
    threads = []
    for body, to_email in emails:
        thread = threading.Thread(
            target=send_email, args=(to_email, subject, username, password, body)
        )
        threads.append(thread)
        thread.start()

    for thread in threads:
        thread.join()


if __name__ == "__main__":
    csv_path = "contacts.csv"
    template_path = "template.docx"
    subject = "Invite to Event in 2023"

    parser = argparse.ArgumentParser(
        description="Send emails using a CSV file and a Word document template."
    )

    parser.add_argument("--email", required=True, help="Your email address.")
    parser.add_argument(
        "--password",
        required=True,
        help="Your email password. If not provided, you'll be prompted to enter it securely.",
    )

    args = parser.parse_args()

    if args.password is None:
        args.password = getpass.getpass(prompt="Enter your email password: ")
    start_time = time.time()

    emails_to_send = [
        (
            fill_invitation(
                template_path,
                {
                    "[Salutation]": row["salutation"],
                    "[First Name]": row["first_name"],
                    "[Last Name]": row["last_name"],
                    "[Last Contacted]": row["last_contacted"],
                    "[Company Name]": row["company"],
                },
            ),
            row["email"],
        )
        for _, row in pd.read_csv(csv_path).iterrows()
    ]

    send_emails(emails_to_send, subject, args.email, args.password)

    end_time = time.time()
    elapsed_time = end_time - start_time
    print(f"Total time taken: {elapsed_time:.2f} seconds")
